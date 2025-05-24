# docx_reader.py

from docx import Document

class DocxReader:
	def __init__(self, file_path):
		self.file_path = file_path
		self.document = Document(file_path)

	def get_all_text(self):
		"""Returns all text in the document as a single string."""
		return "\n".join([para.text for para in self.document.paragraphs])

	def get_paragraphs(self):
		"""Returns a list of all paragraphs."""
		return [para.text for para in self.document.paragraphs]

	def get_tables(self):
		"""Returns all tables as a list of lists (rows of cells)."""
		tables_data = []
		for table in self.document.tables:
			table_data = []
			for row in table.rows:
				row_data = [cell.text.strip() for cell in row.cells]
				table_data.append(row_data)
			tables_data.append(table_data)
		return tables_data

	def get_metadata(self):
		"""Returns document metadata (if available)."""
		props = self.document.core_properties
		return {
			'title': props.title,
			'author': props.author,
			'created': props.created,
			'last_modified_by': props.last_modified_by,
		}

# Example usage (remove or guard under __name__ if using as a module):
if __name__ == "__main__":
	reader = DocxReader("example.docx")
	print("All Text:\n", reader.get_all_text())
	print("\nParagraphs:\n", reader.get_paragraphs())
	print("\nTables:\n", reader.get_tables())
	print("\nMetadata:\n", reader.get_metadata())
